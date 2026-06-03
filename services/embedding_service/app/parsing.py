"""Dosya -> ham metin cikarimi (PDF / DOCX / TXT)."""

from __future__ import annotations

import os
import re


def _normalize_whitespace(text: str) -> str:
    """Bosluklari sadelestir.

    - Windows/Mac satir sonlarini \n yapar.
    - Satir ici fazlalik boslugu tek bosluga indirir.
    - 2'den fazla ardisik bos satiri tek bos satira indirir
      (paragraf sinirlari korunsun diye).
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Her satirin basindaki/sonundaki bosluklari ve satir ici fazla boslugu temizle.
    lines = [re.sub(r"[ \t\f\v]+", " ", line).strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Cok sayida bos satiri tek bos satira indir (paragraf ayraci olarak korunur).
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(data: bytes) -> str:
    """PyMuPDF (fitz) ile PDF metnini cikarir."""
    import fitz  # PyMuPDF; import burada ki modul yuklenirken zorunlu olmasin

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """python-docx ile DOCX paragraf metnini cikarir."""
    import io

    from docx import Document  # python-docx

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Tablolardaki hucre metinlerini de dahil et.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    """UTF-8 (hata toleransli) duz metin cozumu."""
    return data.decode("utf-8", errors="replace")


def extract_text(data: bytes, filename: str) -> str:
    """Dosya icerigini ham metne donusturur.

    Desteklenen uzantilar: .pdf, .docx, .txt.
    Desteklenmeyen uzanti veya bos metin -> ValueError.
    """
    if not data:
        raise ValueError("Bos dosya icerigi.")

    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        raw = _extract_pdf(data)
    elif ext == ".docx":
        raw = _extract_docx(data)
    elif ext == ".txt":
        raw = _extract_txt(data)
    else:
        raise ValueError(f"Desteklenmeyen dosya uzantisi: {ext or '(yok)'}")

    text = _normalize_whitespace(raw)
    if not text:
        raise ValueError("Dosyadan metin cikarilamadi (bos icerik).")
    return text
